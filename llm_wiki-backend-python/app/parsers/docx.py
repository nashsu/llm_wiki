"""DOCX document parser using python-docx."""

import logging
import re
from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from app.parsers.base import DocumentParser, ExtractedImage, ParseResult
from app.parsers.registry import register_parser

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------


def _extract_run_text(run) -> str:
    """Extract text from a *run*, preserving bold / italic as Markdown."""
    text = run.text
    if not text:
        return ""
    if run.bold:
        text = f"**{text}**"
    if run.italic:
        text = f"*{text}*"
    return text


def _paragraph_to_markdown(paragraph) -> str:
    """Convert a ``Paragraph`` to a Markdown line.

    Heading styles become ``#`` / ``##`` / ``###`` … prefixed lines.
    Run-level bold / italic is preserved inline.
    """
    text = paragraph.text.strip()
    if not text:
        return ""

    style_name = paragraph.style.name if paragraph.style else ""

    # Heading styles
    heading_match = re.match(r"^Heading\s+(\d+)$", style_name)
    if heading_match:
        level = int(heading_match.group(1))
        if 1 <= level <= 6:
            prefix = "#" * level
            return f"{prefix} {text}"

    # Title style → level-1 heading
    if style_name == "Title":
        return f"# {text}"

    # Regular paragraph — preserve bold/italic via runs
    parts = [_extract_run_text(run) for run in paragraph.runs]
    return "".join(parts)


def _table_to_markdown(table) -> str:
    """Convert a ``Table`` to a Markdown table string."""
    rows = []
    col_count = 0
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        col_count = max(col_count, len(cells))
        rows.append(cells)

    if not rows:
        return ""

    # Normalise all rows to the same column count
    normalised = []
    for row in rows:
        while len(row) < col_count:
            row.append("")
        normalised.append(row)

    lines: list[str] = []
    # Header row
    lines.append("| " + " | ".join(normalised[0]) + " |")
    # Separator
    lines.append("| " + " | ".join(["---"] * col_count) + " |")
    # Data rows
    for row in normalised[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Parser
# ---------------------------------------------------------------------------


@register_parser()
class DocxParser(DocumentParser):
    """Parser for ``.docx`` documents using python-docx.

    Extracts:
    - Paragraph text with heading → Markdown heading conversion
    - Bold / italic formatting preserved inline
    - Tables converted to Markdown tables
    - Embedded images
    """

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx"]

    @property
    def supports_images(self) -> bool:
        return True

    def parse(self, file_path: str) -> ParseResult:
        path = Path(file_path)
        try:
            doc = DocxDocument(str(path))
        except Exception as exc:
            logger.error("Failed to open DOCX %s: %s", path, exc)
            return ParseResult(
                text="",
                images=[],
                metadata={
                    "source": str(path),
                    "extension": path.suffix,
                    "parser": self.name,
                },
                success=False,
                error=f"Failed to open DOCX: {exc}",
            )

        try:
            text_parts: list[str] = []
            images: list[ExtractedImage] = []

            # Walk the XML body in document order so paragraphs and tables
            # keep their relative position.
            body = doc.element.body
            para_index = 0
            table_index = 0

            for child in body:
                if child.tag == qn("w:p"):
                    if para_index < len(doc.paragraphs):
                        md = _paragraph_to_markdown(doc.paragraphs[para_index])
                        if md:
                            text_parts.append(md)
                        para_index += 1
                elif child.tag == qn("w:tbl"):
                    if table_index < len(doc.tables):
                        md = _table_to_markdown(doc.tables[table_index])
                        if md:
                            text_parts.append(md)
                            text_parts.append("")  # blank line after table
                        table_index += 1

            # Extract embedded images via relationships API
            from docx.opc.constants import RELATIONSHIP_TYPE as RT

            for rel in doc.part.rels.values():
                if rel.reltype == RT.IMAGE:
                    try:
                        image_part = rel.target_part
                        image_data = image_part.blob
                        content_type = image_part.content_type

                        # Derive a meaningful filename
                        part_path = str(image_part.partname) if hasattr(image_part, "partname") else rel.rId
                        import os

                        filename = os.path.basename(part_path) or f"image_{rel.rId}"

                        images.append(
                            ExtractedImage(
                                filename=filename,
                                data=image_data,
                                mime_type=content_type,
                            )
                        )
                    except Exception as img_exc:
                        logger.warning("Failed to extract image %s: %s", rel.rId, img_exc)

            text = "\n\n".join(text_parts)

            metadata: dict = {
                "source": str(path),
                "extension": path.suffix,
                "parser": self.name,
            }

            # Core properties (may fail on some documents)
            try:
                props = doc.core_properties
                if props.title:
                    metadata["title"] = props.title
                if props.author:
                    metadata["author"] = props.author
            except Exception:
                pass

            return ParseResult(text=text, images=images, metadata=metadata)

        except Exception as exc:
            logger.error("Error parsing DOCX %s: %s", path, exc)
            return ParseResult(
                text="",
                images=[],
                metadata={
                    "source": str(path),
                    "extension": path.suffix,
                    "parser": self.name,
                },
                success=False,
                error=f"Error parsing DOCX: {exc}",
            )
